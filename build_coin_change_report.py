from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_DOCX = ROOT / "找零钱问题的多算法求解与适用性分析.docx"
RESULT_IMG = ROOT / "coin_change_result.png"


def set_east_asia_font(run, font_name="宋体"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_text(cell, text, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    set_east_asia_font(run)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "BFBFBF")


def add_paragraph(doc, text="", style=None, first_line=True):
    p = doc.add_paragraph(style=style)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    set_east_asia_font(run)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 77, 120)
    run.font.size = Pt(14 if level == 1 else 12)
    set_east_asia_font(run, "黑体")
    return p


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    cell = table.cell(0, 0)
    shade_cell(cell, "F7F7F7")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(code.strip("\n").splitlines()):
        if i:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)


def create_result_image():
    lines = [
        "实验运行结果示意",
        "硬币面值: [1, 3, 4], 目标金额: 6",
        "贪心算法结果: 4 + 1 + 1, 共 3 枚硬币",
        "动态规划结果: 3 + 3, 共 2 枚硬币",
        "结论: 在该面值体系下，贪心算法未得到最优解。",
        "",
        "硬币面值: [1, 5, 10, 25], 目标金额: 30",
        "贪心算法结果: 25 + 5, 共 2 枚硬币",
        "动态规划结果: 25 + 5, 共 2 枚硬币",
        "结论: 在常见面值体系下，贪心算法可以得到最优解。",
    ]
    width, height = 1100, 520
    image = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    font_paths = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
    ]
    title_font = None
    body_font = None
    for path in font_paths:
        try:
            title_font = ImageFont.truetype(path, 30)
            body_font = ImageFont.truetype(path, 22)
            break
        except OSError:
            continue
    if title_font is None:
        title_font = ImageFont.load_default()
        body_font = ImageFont.load_default()
    draw.rectangle([20, 20, width - 20, height - 20], outline="#8AA4C8", width=3)
    y = 42
    for idx, line in enumerate(lines):
        font = title_font if idx == 0 else body_font
        color = "#1F4D78" if idx == 0 else "#222222"
        draw.text((55, y), line, font=font, fill=color)
        y += 46 if idx == 0 else 40
    image.save(RESULT_IMG)


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)


def add_cover(doc):
    for _ in range(2):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("科技学院\n实验报告")
    r.bold = True
    r.font.size = Pt(26)
    set_east_asia_font(r, "华文行楷")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("(2024-2025年度第2学期)")
    r.font.size = Pt(14)
    set_east_asia_font(r, "宋体")

    doc.add_paragraph()
    rows = [
        ("名    称：", "算法设计与分析"),
        ("院    系：", "计算机系"),
        ("班    级：", ""),
        ("学    号：", ""),
        ("学生姓名：", ""),
        ("成    绩：", ""),
        ("日    期：", "2025年 6 月 20 日"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows):
        set_cell_text(table.cell(i, 0), label, bold=True, size=12, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(table.cell(i, 1), value, size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
    for row in table.rows:
        row.height = Cm(1.0)
    doc.add_page_break()


def build_doc():
    create_result_image()
    doc = Document()
    setup_styles(doc)
    add_cover(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("找零钱问题的多算法求解与适用性分析")
    run.bold = True
    run.font.size = Pt(18)
    set_east_asia_font(run, "黑体")

    add_heading(doc, "一、问题描述", 1)
    add_paragraph(
        doc,
        "找零钱问题是算法设计中的经典组合优化问题。给定若干种不同面值的硬币，以及一个目标金额，要求选择若干枚硬币，使它们的面值之和等于目标金额，并且所用硬币数量尽可能少。"
    )
    add_paragraph(
        doc,
        "例如硬币面值为 {1, 3, 4}，目标金额为 6。若采用 3+3，只需要 2 枚硬币；若采用 4+1+1，则需要 3 枚硬币。因此，该问题不仅要求找到可行解，还要求在所有可行方案中找到最优解。"
    )

    add_heading(doc, "二、算法思想", 1)
    add_heading(doc, "1. 递归法", 2)
    add_paragraph(
        doc,
        "递归法从目标金额 amount 出发，尝试选择每一种不超过当前金额的硬币，然后递归求解剩余金额 amount-coin。递归边界是 amount=0，此时不需要硬币；若 amount<0，则说明该路径不可行。递归法思想直观，但会重复计算大量相同子问题。"
    )

    add_heading(doc, "2. 回溯法", 2)
    add_paragraph(
        doc,
        "回溯法把每一次选择硬币看成搜索树中的一层，通过深度优先搜索枚举不同硬币组合。当当前金额超过目标金额，或当前硬币数已经不小于已知最优解时，停止继续搜索。回溯法能够找到最优解，并且可以通过剪枝减少搜索量，但最坏情况下仍可能接近指数级复杂度。"
    )

    add_heading(doc, "3. 贪心算法", 2)
    add_paragraph(
        doc,
        "贪心算法每一步都选择不超过剩余金额的最大面值硬币，直到凑够目标金额。它的优点是实现简单、速度快；但它只关注当前局部最优选择，不能保证在所有硬币面值体系下都得到全局最优解。"
    )

    add_heading(doc, "4. 动态规划法", 2)
    add_paragraph(
        doc,
        "动态规划法利用最优子结构。设 dp[i] 表示凑出金额 i 所需要的最少硬币数，则状态转移方程为：dp[i] = min(dp[i], dp[i-coin] + 1)，其中 coin 为某一种硬币面值且 i>=coin。初始状态 dp[0]=0，其余状态设为无穷大。动态规划避免了重复计算，能够稳定求出最优解。"
    )

    add_heading(doc, "三、程序实现", 1)
    add_paragraph(doc, "以下给出四种方法的关键代码，省略输入输出等非核心部分。", first_line=False)

    add_heading(doc, "1. 递归法关键代码", 2)
    add_code_block(
        doc,
        """
def coin_change_recursive(coins, amount):
    if amount == 0:
        return 0
    if amount < 0:
        return float('inf')
    ans = float('inf')
    for coin in coins:
        ans = min(ans, coin_change_recursive(coins, amount - coin) + 1)
    return ans
""",
    )
    add_paragraph(doc, "该代码通过递归枚举当前可选择的硬币，并把问题转化为剩余金额的最少硬币数问题。")

    add_heading(doc, "2. 回溯法关键代码", 2)
    add_code_block(
        doc,
        """
def coin_change_backtrack(coins, amount):
    coins.sort(reverse=True)
    best = [float('inf')]

    def dfs(start, remain, count):
        if remain == 0:
            best[0] = min(best[0], count)
            return
        if count >= best[0]:
            return
        for i in range(start, len(coins)):
            if coins[i] <= remain:
                dfs(i, remain - coins[i], count + 1)

    dfs(0, amount, 0)
    return best[0]
""",
    )
    add_paragraph(doc, "该代码按照硬币面值从大到小搜索，并用当前硬币数与已知最优解比较进行剪枝，从而减少不必要的搜索。")

    add_heading(doc, "3. 贪心算法关键代码", 2)
    add_code_block(
        doc,
        """
def coin_change_greedy(coins, amount):
    coins.sort(reverse=True)
    result = []
    for coin in coins:
        while amount >= coin:
            amount -= coin
            result.append(coin)
    return result if amount == 0 else None
""",
    )
    add_paragraph(doc, "该代码每次选择当前可用的最大面值硬币，直到目标金额被凑出或无法继续选择。")

    add_heading(doc, "4. 动态规划法关键代码", 2)
    add_code_block(
        doc,
        """
def coin_change_dp(coins, amount):
    dp = [float('inf')] * (amount + 1)
    dp[0] = 0
    for i in range(1, amount + 1):
        for coin in coins:
            if i >= coin:
                dp[i] = min(dp[i], dp[i - coin] + 1)
    return dp[amount]
""",
    )
    add_paragraph(doc, "该代码自底向上计算每一个金额的最优解，最终 dp[amount] 即为目标金额所需的最少硬币数。")

    add_heading(doc, "四、实验结果与分析", 1)
    add_paragraph(doc, "为比较不同算法的适用性，选取两组硬币面值进行测试。第一组 {1, 3, 4} 用于说明贪心算法可能失效，第二组 {1, 5, 10, 25} 用于说明贪心算法在常见面值体系下表现较好。")
    doc.add_picture(str(RESULT_IMG), width=Cm(15.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    headers = ["算法", "基本思想", "是否保证最优", "时间复杂度", "适用场景"]
    for i, header in enumerate(headers):
        set_cell_text(table.cell(0, i), header, bold=True)
        shade_cell(table.cell(0, i), "E8EEF5")
    rows = [
        ["递归法", "枚举每种选择并递归求解子问题", "是", "指数级", "问题规模较小、用于理解问题结构"],
        ["回溯法", "搜索所有可能组合并进行剪枝", "是", "最坏指数级", "需要输出方案或规模中等时"],
        ["贪心算法", "每次选当前最大可用面值", "不一定", "O(k + m)", "满足贪心性质的硬币体系"],
        ["动态规划", "保存子问题最优解并递推", "是", "O(nk)", "通用求最优解场景"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT if i in (1, 4) else WD_ALIGN_PARAGRAPH.CENTER)

    add_paragraph(
        doc,
        "其中 n 表示目标金额，k 表示硬币种类数，m 表示最终选择的硬币枚数。递归法和回溯法都具有枚举性质，适合展示问题的搜索过程，但当金额较大时效率较低。贪心算法速度最快，但只有在硬币面值满足特定性质时才能保证最优。动态规划法时间复杂度为 O(nk)，虽然需要额外数组空间，但适用范围最广，能够保证得到最优解。"
    )
    add_paragraph(
        doc,
        "通过 {1, 3, 4}、amount=6 的例子可以看出，贪心算法选择 4+1+1，需要 3 枚硬币；而最优解是 3+3，只需要 2 枚硬币。因此，找零钱问题不能简单地依赖贪心策略。动态规划法通过系统比较所有子问题的最优解，可以避免局部最优导致整体非最优的问题。"
    )

    add_heading(doc, "五、课程设计总结", 1)
    add_paragraph(
        doc,
        "通过本次专题报告，我进一步理解了同一个问题可以从递归、回溯、贪心和动态规划等不同角度进行求解。递归强调问题分解，回溯强调搜索与剪枝，贪心强调局部选择，动态规划则强调状态表示和最优子结构。"
    )
    add_paragraph(
        doc,
        "找零钱问题表面上比较简单，但它很好地体现了算法适用性的差异。尤其是贪心算法并不总能得到最优解，这说明在设计算法时不能只关注代码是否简洁，还要分析算法成立的条件。通过与动态规划法对比，可以看出动态规划虽然实现稍复杂，但在求解最优问题时更稳定可靠。"
    )
    add_paragraph(
        doc,
        "本课程使我认识到算法设计不仅是写出程序，更重要的是分析问题结构、选择合适策略，并从时间复杂度、空间复杂度和适用条件等方面评价算法。建议课程中可以增加更多生活化案例和算法对比实验，这样有助于理解不同算法思想之间的联系与区别。"
    )

    doc.save(OUT_DOCX)


if __name__ == "__main__":
    build_doc()
    print(OUT_DOCX)
